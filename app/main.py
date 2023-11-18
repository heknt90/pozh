import os
import json
from docx import Document

src_path = "./src/"
dirs = ['commandersData', 'commandersDataWithGDZS', 'driversData', 'firemansData', 'firemansDataWithGDZS']
result_dict = {}
for dir in dirs:
  result_dict[dir] = {}
  path = src_path + dir + "/"
  length = len(os.listdir(path))
  for i in range(1, length):
    ticket_dict = {}
    filename = path + str(i) + ".docx"
    document = Document(filename)
    tables = document.tables
    for table in tables:
      def check_answer(cells):
        for p in cells[2].paragraphs:
          for r in p.runs:
            if str(r.font.color.rgb) == "FF0000":
              return True
            
      def get_qa(cells):
        key = cells[0].text
        value = {
          "q": cells[1].text,
          "a": cells[2].text
        }
        return (key, value)

      for row in table.rows:
        cells = row.cells
        if check_answer(cells):
          qa = get_qa(cells)
          ticket_dict[qa[0]] = qa[1]
      # print(ticket_dict)
      # print("=============")
      # print()
    result_dict[dir][i] = ticket_dict 
# print(result_dict)
# data = json.dumps(result_dict, separators=(',', ':'))
# data = json.dumps(result_dict)
# print(data)
with open("data.json", "w", encoding="utf-8") as f:
  json.dump(result_dict, f, ensure_ascii=False, indent=4)
# f = open("data.json", "wb")
# f.write(data)